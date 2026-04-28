"""
Generate two Word manuals — Field Manual and Admin Manual — for the
BAK Construction / Bengal Crane and Rigging management system.

Branding: Bengal Crane red (#B91C1C) accent + black headings, white body.
If a logo file exists at ./bengal-crane-logo.png it gets dropped into the
cover page; otherwise the cover shows a typographic stand-in.

Run: python3 build_manuals.py
Outputs: BAK-Field-Manual.docx, BAK-Admin-Manual.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
LOGO = HERE / "bengal-crane-logo.png"

# Bengal Crane red — adjust if Brenda gives an exact hex.
BRAND_RED = RGBColor(0xB9, 0x1C, 0x1C)
BLACK = RGBColor(0x11, 0x11, 0x11)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0xE5, 0xE5, 0xE5)


def shade(cell, hex_color):
    """Set a table cell's background shading."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="BBBBBB", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = BRAND_RED
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "B91C1C")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = BLACK
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = BRAND_RED
    return p


def add_para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = ""
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    if p.runs:
        p.runs[0].text = ""
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    return p


def add_callout(doc, label, body):
    """Single-cell shaded callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "FEF2F2")  # soft red wash
    set_cell_borders(cell, color="FCA5A5", size="6")
    p1 = cell.paragraphs[0]
    r = p1.add_run(label.upper())
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BRAND_RED
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(11)
    doc.add_paragraph()  # spacer


def add_step_table(doc, steps):
    """Two-column 'step / what to do' table."""
    table = doc.add_table(rows=len(steps) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(0.7)
    table.columns[1].width = Inches(5.8)

    # Header
    hdr_a = table.cell(0, 0)
    hdr_b = table.cell(0, 1)
    shade(hdr_a, "111111")
    shade(hdr_b, "111111")
    for c, text in [(hdr_a, "STEP"), (hdr_b, "WHAT TO DO")]:
        set_cell_borders(c, color="111111")
        p = c.paragraphs[0]
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, body in enumerate(steps, start=1):
        a = table.cell(i, 0)
        b = table.cell(i, 1)
        set_cell_borders(a)
        set_cell_borders(b)
        pa = a.paragraphs[0]
        ra = pa.add_run(str(i))
        ra.font.bold = True
        ra.font.size = Pt(12)
        ra.font.color.rgb = BRAND_RED
        pa.alignment = WD_ALIGN_PARAGRAPH.CENTER

        pb = b.paragraphs[0]
        rb = pb.add_run(body)
        rb.font.name = "Calibri"
        rb.font.size = Pt(11)

    doc.add_paragraph()


def add_cover(doc, title, subtitle, audience):
    """Cover page with branding strip + title block."""
    # Top brand strip (single-cell table acts as a colored band)
    band = doc.add_table(rows=1, cols=1)
    cell = band.cell(0, 0)
    shade(cell, "B91C1C")
    p = cell.paragraphs[0]
    r = p.add_run("BENGAL CRANE & RIGGING  ·  BAK CONSTRUCTION")
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.bold = True
    r.font.size = Pt(13)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph()

    # Logo or stand-in
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(2.4))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("[ LOGO ]")
        r.font.name = "Calibri"
        r.font.size = Pt(28)
        r.font.bold = True
        r.font.color.rgb = LIGHT_GRAY

    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title.upper())
    r.font.name = "Calibri"
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = BLACK

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    r.font.color.rgb = GRAY
    r.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Audience block
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Inches(4.5)
    bcell = box.cell(0, 0)
    shade(bcell, "F3F4F6")
    set_cell_borders(bcell, color="111111", size="8")
    bp = bcell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run("INTENDED AUDIENCE")
    br.font.bold = True
    br.font.size = Pt(10)
    br.font.color.rgb = BRAND_RED
    bp2 = bcell.add_paragraph()
    bp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br2 = bp2.add_run(audience)
    br2.font.size = Pt(13)
    br2.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer/version
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Version 1.0  ·  Draft for review")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    doc.add_page_break()


def add_toc_placeholder(doc, items):
    add_h1(doc, "Table of Contents")
    for i, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}.  {item}")
        r.font.name = "Calibri"
        r.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


def set_default_font(doc, name="Calibri", size=11):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)


# ──────────────────────────────────────────────────────────────────────
# FIELD MANUAL
# ──────────────────────────────────────────────────────────────────────
def build_field_manual():
    doc = Document()
    set_default_font(doc)

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    add_cover(
        doc,
        "Field Manual",
        "Daily-use guide for foremen and field crew",
        "Foremen, Site Managers, and Field Crew",
    )

    add_toc_placeholder(doc, [
        "Getting Started — logging in",
        "Clocking In and Out",
        "Filling out a Timesheet (Single Entry)",
        "Filling out Timesheets in Bulk (Crew or Batch)",
        "Daily Logs",
        "Uploading Job Photos",
        "Equipment Check-In / Check-Out",
        "Submitting for Approval",
        "Common Issues and How to Fix Them",
    ])

    # ── 1. Getting Started ──
    add_h1(doc, "1. Getting Started")
    add_para(doc, "Welcome to the BAK Construction management system. This manual covers everything a foreman or field worker needs to do day-to-day. The office team uses a separate Admin Manual for setup and approvals.")
    add_h2(doc, "Logging In")
    add_step_table(doc, [
        "Open your web browser (Chrome or Safari work best) and go to the address your office sent you.",
        "Type your email and password. If you forgot your password, click \"Forgot Password\" or ask the office to reset it.",
        "Click Sign In. You'll land on the dashboard with shortcut tiles.",
    ])
    add_callout(doc, "Tip",
        "Bookmark the login page on your phone's home screen so you can open it with one tap from the jobsite.")

    # ── 2. Clocking In/Out ──
    add_h1(doc, "2. Clocking In and Out")
    add_para(doc, "Two ways to track time on a job:")
    add_h2(doc, "A. My Time (each worker clocks themselves)")
    add_step_table(doc, [
        "From the dashboard, tap My Time.",
        "Pick the project you're working on from the dropdown.",
        "Tap Clock In. The system stamps the time and your GPS location.",
        "When you take an unpaid lunch, tap Clock Out, then Clock In again when you're back.",
        "At the end of the day, tap Clock Out one last time.",
    ])
    add_h2(doc, "B. Crew Clock (foreman clocks the whole crew at once)")
    add_step_table(doc, [
        "From the dashboard, tap Crew Clock-In.",
        "Pick the project.",
        "Tick the workers who are on-site today.",
        "Tap Clock In Crew — every selected worker gets stamped at once.",
        "At the end of the day, repeat with Crew Clock-Out.",
    ])
    add_callout(doc, "Important",
        "If GPS isn't working (basement, deep inside a building), the system still records the time — just not the location. Note it in the comments so the office knows.")

    # ── 3. Single Timesheet ──
    add_h1(doc, "3. Filling Out a Timesheet (Single Entry)")
    add_para(doc, "Use this when you need to add or correct one entry — for example, a worker forgot to clock in.")
    add_step_table(doc, [
        "Go to Timesheets > Add Timesheet.",
        "Pick the date, employee, and project.",
        "Enter Hours Worked (the system splits ST and OT automatically using the weekly 40-hour rule).",
        "If the work was OT regardless (holiday, weekend premium), tick Force OT.",
        "Pick the Earnings Category: HE for regular worked hours, HO for Holiday pay, or VA for Vacation.",
        "Tick Per Diem if it applies. Add notes if anything is unusual.",
        "Click Save.",
    ])

    # ── 4. Bulk Entry ──
    add_h1(doc, "4. Filling Out Timesheets in Bulk")
    add_para(doc, "Best when you're keying in a full day or week for multiple workers. There are two layouts to choose from:")
    add_h2(doc, "Crew Roster (Mark Whole Crew Present/Absent)")
    add_para(doc, "Pick a crew, set the date and shift, and you'll see every member listed with one row each. Tick Present, type hours, click Save.")
    add_h2(doc, "Batch Entry (Foundation-style payroll keying)")
    add_step_table(doc, [
        "Go to Timesheets > Bulk Entry.",
        "Pick the time period (begin / end / W-E date) at the top.",
        "Type or pick the Job, Work Date, Employee #, Cost Code, Cost Type, and Craft.",
        "Enter ST / OT / PR hours and the Earnings Category.",
        "Press F10 (or click Save Record) — the entry posts and clears for the next one. Job, Date, and Shift carry over.",
        "Repeat for each line. The running list at the bottom shows everything you've keyed in this batch.",
    ])
    add_callout(doc, "Holiday and Vacation",
        "When you set Earnings Category to HO (Holiday) or VA (Vacation), the OT and PR hour fields lock — those pay types are flat rate only.")

    # ── 5. Daily Logs ──
    add_h1(doc, "5. Daily Logs")
    add_para(doc, "Daily logs capture what happened on site that day — weather, work performed, manpower, deliveries, visitors. Required at the end of every shift.")
    add_step_table(doc, [
        "Go to Daily Logs > Add Daily Log.",
        "Pick the date and project.",
        "Fill in: weather, temperature, work performed, manpower count, equipment used.",
        "Note anything unusual: delays, accidents, visitors, deliveries, RFIs.",
        "Attach photos (see next section) if needed.",
        "Click Submit. The Site Manager will review.",
    ])

    # ── 6. Photos ──
    add_h1(doc, "6. Uploading Job Photos")
    add_step_table(doc, [
        "From the project page, tap Photos > Upload.",
        "Select photos from your phone gallery (you can pick multiple at once).",
        "Add a quick caption — what it shows, what date, anything notable.",
        "Tap Upload.",
    ])
    add_callout(doc, "Tip",
        "Take photos every morning, every evening, and any time something changes (deliveries, damage, completion). The office uses these for billing back-up and dispute resolution.")

    # ── 7. Equipment ──
    add_h1(doc, "7. Equipment Check-In / Check-Out")
    add_para(doc, "Every piece of equipment that arrives on-site or leaves needs to be recorded so the rental calendar stays accurate.")
    add_h3(doc, "When equipment arrives")
    add_step_table(doc, [
        "Go to the project > Equipment tab.",
        "Click Assign Equipment.",
        "Pick the piece of equipment (or the vendor's record if it's a rental).",
        "Enter the Assigned Date (today) and the Expected Return Date.",
        "Click Save. The piece now appears on the Rental Calendar.",
    ])
    add_h3(doc, "When equipment leaves")
    add_step_table(doc, [
        "Find the assignment on the project's Equipment tab.",
        "Click Mark Returned and enter the actual return date.",
        "Save.",
    ])

    # ── 8. Submitting ──
    add_h1(doc, "8. Submitting for Approval")
    add_para(doc, "Timesheets you create start in Draft. To send them to the Site Manager or Brenda for sign-off:")
    add_step_table(doc, [
        "Go to the Timesheets list.",
        "Tick the rows you want to submit, OR open one and click Submit.",
        "The status changes to Submitted. The Site Manager will Approve or Reject.",
        "Approved timesheets are locked — no more edits without an admin's help.",
        "Rejected timesheets come back to Draft with a note explaining why. Fix and re-submit.",
    ])

    # ── 9. Troubleshooting ──
    add_h1(doc, "9. Common Issues and How to Fix Them")
    issues = [
        ("Can't log in", "Double-check your email and password. Try Forgot Password. If still stuck, ask the office to reset it."),
        ("My hours look wrong (split into OT incorrectly)", "Open the timesheet, tick Force OT to override the weekly-40 rule, and save. Or contact the office."),
        ("Can't find a project in the dropdown", "Project may be set to inactive or hidden. Ask the office to make it active."),
        ("Photo upload stuck or failing", "Try a smaller photo (under 10MB). Switch to Wi-Fi if you're on cellular."),
        ("Approve button missing", "Only Site Managers and Admins can approve. Foremen submit; the Site Manager signs off."),
    ]
    for q, a in issues:
        add_h3(doc, q)
        add_para(doc, a)

    add_h1(doc, "Questions?")
    add_para(doc, "Office contact: Brenda — phone, email, and office hours go here once finalized.", italic=True, color=GRAY)

    return doc


# ──────────────────────────────────────────────────────────────────────
# ADMIN MANUAL
# ──────────────────────────────────────────────────────────────────────
def build_admin_manual():
    doc = Document()
    set_default_font(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    add_cover(
        doc,
        "Admin Manual",
        "Office-side guide for setup, approvals, billing, and reports",
        "Office Staff, Accountants, Site Managers, and Brenda",
    )

    add_toc_placeholder(doc, [
        "System Overview & Roles",
        "User Accounts and Permissions",
        "Setting Up Projects",
        "Employees, Crafts, and Crews",
        "Vendors, Materials, and Equipment",
        "Estimates → Awarded Projects",
        "Change Orders",
        "Purchase Orders & Receiving",
        "Timesheets — Review and Approval",
        "Payroll Exports",
        "Invoicing & Billing",
        "Reports — Job Costing, WIP, Equipment Utilization",
        "Database Backups",
        "Settings (Logo, Company Info, Rates)",
    ])

    # ── 1. Overview ──
    add_h1(doc, "1. System Overview & Roles")
    add_para(doc, "The BAK Construction management system covers the full lifecycle: estimating, awarded projects, daily field operations, purchasing, labor tracking, billing, and financial reporting.")

    add_h2(doc, "Roles and Who Does What")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, t in enumerate(["ROLE", "WHAT THEY CAN DO"]):
        shade(hdr[i], "111111")
        set_cell_borders(hdr[i], color="111111")
        p = hdr[i].paragraphs[0]
        r = p.add_run(t)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    role_rows = [
        ("Admin (Brenda)", "Full access. Approve timesheets, manage users, settings, backups."),
        ("Site Manager", "Approve/reject timesheets. View all projects, daily logs, photos."),
        ("Project Manager", "Create + manage projects, estimates, change orders, POs. Cannot approve timesheets."),
        ("Foreman", "Create timesheets for crew, daily logs, photos, equipment check-in. Cannot approve."),
        ("Field Staff", "Clock self in/out, fill own timesheet, view assigned project info."),
        ("Accountant", "View all financial pages, run payroll exports, run reports. Cannot approve timesheets."),
        ("Viewer", "Read-only access. For owner/PM-side stakeholders."),
    ]
    for role, desc in role_rows:
        row = table.add_row().cells
        set_cell_borders(row[0])
        set_cell_borders(row[1])
        p = row[0].paragraphs[0]
        r = p.add_run(role)
        r.font.bold = True
        r.font.color.rgb = BRAND_RED
        r.font.size = Pt(11)
        row[1].paragraphs[0].add_run(desc).font.size = Pt(11)

    doc.add_paragraph()

    # ── 2. Users ──
    add_h1(doc, "2. User Accounts and Permissions")
    add_h2(doc, "Adding a New User")
    add_step_table(doc, [
        "Go to Settings > Users.",
        "Click Add User.",
        "Enter name, email, password, and role.",
        "Set Active. Click Save.",
        "Tell the user their email + temporary password and ask them to change it on first login.",
    ])
    add_h2(doc, "Deactivating a User")
    add_para(doc, "Don't delete users — uncheck Active instead. That way historic timesheets and approvals keep their author's name on the audit trail.")

    # ── 3. Projects ──
    add_h1(doc, "3. Setting Up Projects")
    add_step_table(doc, [
        "Go to Projects > Add Project.",
        "Fill in: project number (auto-generates if blank), name, client, address, GC, status.",
        "Set Original Budget, Contract Value, and Retainage % if known.",
        "Optional: Default Per Diem Rate, Default Billable Rates.",
        "Click Save. The project shows up on dropdowns everywhere it's needed.",
    ])
    add_callout(doc, "Status meanings",
        "Bidding: estimate stage. Awarded: contract signed. Active: work in progress. Closed: done. Cancelled: dead deal.")

    # ── 4. Employees / Crews ──
    add_h1(doc, "4. Employees, Crafts, and Crews")
    add_h2(doc, "Adding an Employee")
    add_step_table(doc, [
        "Go to Employees > Add Employee.",
        "Enter name, employee number (auto if blank), craft, hourly rate, status.",
        "Optional: Default Cost Type, certifications (under the Certifications tab).",
        "Save.",
    ])
    add_h2(doc, "Importing Employees from Excel")
    add_para(doc, "If you have an existing spreadsheet, click Import and pick the file. The system accepts .xlsx and .csv. Column headers can be lowercase or with underscores — the importer handles both.")
    add_h2(doc, "Crews")
    add_para(doc, "A crew is a group of employees usually deployed together (e.g., \"Crane Crew #2\"). Set a foreman, add members, assign to a project. Bulk timesheet entry uses crews to load the whole roster in one click.")

    # ── 5. Vendors / Equipment ──
    add_h1(doc, "5. Vendors, Materials, and Equipment")
    add_para(doc, "Vendors are anyone you write a check to that isn't payroll — material suppliers, equipment rentals, subcontractors. Each Vendor has POs, payments, and 1099 status.")
    add_h2(doc, "Equipment")
    add_step_table(doc, [
        "Go to Equipment > Add Equipment.",
        "Enter name, type (Owned / Rented / Third-Party), description.",
        "Set daily, weekly, monthly rates if applicable.",
        "Pick the vendor if it's rented.",
        "Save.",
    ])
    add_para(doc, "Equipment shows on the Rental Calendar based on assignments — when you assign a piece to a project with a start date and expected return date, it appears as a Gantt-style bar.")

    # ── 6. Estimates ──
    add_h1(doc, "6. Estimates → Awarded Projects")
    add_step_table(doc, [
        "Go to Estimates > New Estimate.",
        "Pick the client and project type.",
        "Build line items by section (mobilization, labor, equipment, materials, subs, etc).",
        "Set markup, contingency, taxes.",
        "Use the Print/PDF button to send the proposal.",
        "When awarded, click Convert to Project — the line items become the project's budget by cost code.",
    ])

    # ── 7. Change Orders ──
    add_h1(doc, "7. Change Orders")
    add_step_table(doc, [
        "From the project page, click Change Orders > New Change Order.",
        "Enter description, scope of change, line items.",
        "Set status to Pending until the client signs. Once signed, change to Approved — the contract value updates automatically.",
        "Print/email the CO from the show page.",
    ])

    # ── 8. POs ──
    add_h1(doc, "8. Purchase Orders & Receiving")
    add_step_table(doc, [
        "Go to Purchase Orders > New PO.",
        "Pick vendor, project, cost code, and cost type.",
        "Add line items: material, quantity, unit cost.",
        "Save and Send to vendor (PDF or email).",
        "When materials arrive: open the PO, mark items received with quantities. Partial receipts are allowed.",
    ])

    # ── 9. Timesheet review ──
    add_h1(doc, "9. Timesheets — Review and Approval")
    add_para(doc, "Foremen submit timesheets daily. The Site Manager and Brenda review and approve.")
    add_h2(doc, "Daily Review")
    add_step_table(doc, [
        "Go to Timesheets.",
        "Filter by date and status = Submitted.",
        "Open each row, verify hours match the daily log, click Approve.",
        "Or use the bulk approve checkbox — select multiple submitted rows, click Approve Selected.",
    ])
    add_h2(doc, "Rejecting a Timesheet")
    add_para(doc, "Click Reject and add a note explaining why. The foreman gets it back as Draft and can fix and re-submit.")

    add_h2(doc, "Print for Billing")
    add_para(doc, "Click Print for Billing and choose:")
    add_bullet(doc, "Per Timesheet (Daily) — one page per entry, for daily client sign-off.")
    add_bullet(doc, "Weekly Summary — one landscape page per employee per week, with Mon–Sun laid out side by side and weekly totals. Best for billing & payroll.")

    # ── 10. Payroll ──
    add_h1(doc, "10. Payroll Exports")
    add_para(doc, "End of pay period:")
    add_step_table(doc, [
        "Go to Reports > Payroll Export.",
        "Pick the pay period (W/E date).",
        "Pick the export format (CSV or your payroll provider's layout).",
        "Click Download. Hand it to your payroll service or import directly.",
    ])

    # ── 11. Billing ──
    add_h1(doc, "11. Invoicing & Billing")
    add_step_table(doc, [
        "Go to Billing > New Invoice.",
        "Pick the project and the period.",
        "The system pulls in approved timesheets, change orders, and reimbursables for that window.",
        "Adjust as needed (reorder line items, edit descriptions).",
        "Print/PDF and send to the client.",
        "Mark Paid when payment arrives — the receivable on the project P&L closes out.",
    ])

    # ── 12. Reports ──
    add_h1(doc, "12. Reports")
    add_para(doc, "Available reports:")
    add_bullet(doc, "Job Costing — by project, by cost code, by cost type. Shows budget vs. actual.")
    add_bullet(doc, "WIP (Work In Progress) — earned revenue, billed, over/underbilling.")
    add_bullet(doc, "Equipment Utilization — hours used per piece, billable amount, idle time.")
    add_bullet(doc, "Labor by Project / Employee — by week, month, or any custom range.")
    add_bullet(doc, "Certified Payroll — for prevailing wage jobs, in DOL format.")

    # ── 13. Backups ──
    add_h1(doc, "13. Database Backups")
    add_para(doc, "The system runs an automated nightly backup. You can also trigger one manually:")
    add_step_table(doc, [
        "Go to Settings > Backups.",
        "Click Create Backup Now.",
        "When done, the backup appears in the list. Click Download to save a copy off-site.",
    ])
    add_callout(doc, "Recommended",
        "Download a backup before any major data import or schema change. Keep at least 30 days of monthly backups in a separate location.")

    # ── 14. Settings ──
    add_h1(doc, "14. Settings (Logo, Company Info, Rates)")
    add_h2(doc, "Branding")
    add_step_table(doc, [
        "Go to Settings > Branding.",
        "Upload the company logo (PNG with transparent background recommended).",
        "Set the primary color (hex code, e.g. #B91C1C for Bengal Crane red).",
        "Save. The logo and color show on every printout and PDF.",
    ])
    add_h2(doc, "Company Info")
    add_para(doc, "Company name, address, EIN, phone — these appear on invoices, POs, and proposals.")
    add_h2(doc, "Default Rates")
    add_para(doc, "Set the default billable markup, OT multiplier (typically 1.5x), DT multiplier (2x), and per-diem rate. Projects can override these individually.")

    add_h1(doc, "Questions?")
    add_para(doc, "Send revisions, corrections, or additions to the development team — this manual is meant to grow with the system.", italic=True, color=GRAY)

    return doc


def main():
    field = build_field_manual()
    field_path = HERE / "BAK-Field-Manual.docx"
    field.save(str(field_path))
    print(f"✓ {field_path.name} ({field_path.stat().st_size:,} bytes)")

    admin = build_admin_manual()
    admin_path = HERE / "BAK-Admin-Manual.docx"
    admin.save(str(admin_path))
    print(f"✓ {admin_path.name} ({admin_path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
